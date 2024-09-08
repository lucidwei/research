# coding=gbk
# Time Created: 2024/4/3 14:27
# Author  : Lucid
# FileName: run_boom.py
# Software: PyCharm
from base_config import BaseConfig
from prj_boom.preprocess import DataPreprocessor
from prj_boom.modeler import DynamicFactorModeler
from docx import Document
from docx.shared import Inches
import openpyxl
from datetime import datetime
import warnings
import pandas as pd

# 忽略特定的 FutureWarning
warnings.filterwarnings("ignore", category=FutureWarning,
                        message='verbose is deprecated since functions should not print results')

base_config = BaseConfig('boom')

# SINGLE_BATCH_MODE = 'single'
SINGLE_BATCH_MODE = 'batch'
# SINGLE_BATCH_MODE = 'batch_macro'

if SINGLE_BATCH_MODE == 'single':
    # 字典存储不同的配置，每个配置包含industry和compare_to等信息
    # key包括：industry, compare_to, stationary, date_start, single_line
    configs = [
        # {'industry': '石油石化', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '石油石化', 'compare_to': '归属母公司股东的净利润同比增长率', 'stationary': False},
        # {'industry': '煤炭', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '有色金属', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '钢铁', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '基础化工', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '建材', 'compare_to': '净资产收益率ROE', 'stationary': False},
        {'industry': '建筑', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '电子', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '传媒', 'compare_to': '净资产收益率ROE'},
        # {'industry': '计算机', 'compare_to': '净资产收益率ROE'},
        # {'industry': '通信', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '机械', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '电新', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '电力', 'compare_to': '净资产收益率ROE'},
        # {'industry': '国防军工', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '交通运输', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '汽车', 'compare_to': '净资产收益率ROE'},
        # {'industry': '家电', 'compare_to': '净资产收益率ROE'},
        # {'industry': '医药', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '食品饮料', 'compare_to': '净资产收益率ROE'},
        # {'industry': '农林牧渔', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '白酒', 'compare_to': '净资产收益率ROE'},
        # {'industry': '食品', 'compare_to': '净资产收益率ROE'},
        # {'industry': '饮料', 'compare_to': '净资产收益率ROE'},
        # {'industry': '消费服务', 'compare_to': '净资产收益率ROE'},
        # {'industry': '纺服', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '商贸零售', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '美容护理', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '房地产', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '非银金融', 'compare_to': '净资产收益率ROE', 'stationary': False},
        # {'industry': '就业状况', 'compare_to': '中国:城镇调查失业率', 'stationary': False, 'date_start': '2018-01-01'},
        # {'industry': '社零综指', 'compare_to': '中国:社会消费品零售总额:当月同比', 'stationary': False, 'date_start': '2020-01-01'},
        {'industry': '出口', 'compare_to': '中国:出口金额:当月同比', 'stationary': False, 'date_start': '2010-01-01'},
        # {'industry': '出口', 'compare_to': '中国:出口金额:当月同比', 'stationary': False, 'date_start': '2010-01-01', 'single_line': True},
        # {'industry': 'PPI', 'compare_to': '中国:PPI:全部工业品:当月同比', 'stationary': False, 'date_start': '2010-01-01', 'single_line': True},
        # {'industry': '房价', 'compare_to': '中国:房屋销售价格指数:二手住宅:70个大中城市:当月同比', 'stationary': False, 'date_start': '2010-01-01'},
        # {'industry': '工业增加值', 'compare_to': '中国:规模以上工业增加值:当月同比', 'stationary': False, 'date_start': '2010-01-01'},
        # {'industry': '制造业投资', 'compare_to': '(月度化)中国:固定资产投资完成额:制造业:累计同比', 'stationary': False, 'date_start': '2014-01-01', 'leading_prediction': False},
        # {'industry': '制造业投资', 'compare_to': '中国:固定资产投资完成额:制造业:累计同比', 'stationary': False, 'date_start': '2014-01-01', 'leading_prediction': False},
        # {'industry': '房地产投资', 'compare_to': '中国:房地产开发投资完成额:累计同比', 'stationary': False, 'date_start': '2012-01-01', 'leading_prediction': False},
        # {'industry': '房地产投资', 'compare_to': '(月度化)中国:房地产开发投资完成额:累计同比', 'stationary': False, 'date_start': '2012-01-01', 'leading_prediction': False, 'plot_y0': True},
        # {'industry': '基建投资', 'compare_to': '(月度化)中国:固定资产投资完成额:基础设施建设投资:累计同比', 'stationary': False, 'date_start': '2012-01-01', 'leading_prediction': False, 'plot_y0': True},
        # {'industry': '基建投资', 'compare_to': '(月度化)中国:固定资产投资本年新开工项目计划总投资额:累计同比', 'stationary': False, 'date_start': '2012-01-01', 'leading_prediction': False, 'plot_y0': True},
        # {'industry': '基建投资', 'compare_to': '(月度化)中国:固定资产投资本年施工项目计划总投资额:累计同比', 'stationary': False, 'date_start': '2012-01-01', 'leading_prediction': False, 'plot_y0': True},
    ]

    for config in configs:
        preprocessor = DataPreprocessor(
            base_config,
            industry=config['industry'],
            stationary=config.get('stationary', False),
            date_start=config.get('date_start', '2013-01-01'),
            compare_to=config['compare_to']
        )
        preprocessor.preprocess()

        modeler = DynamicFactorModeler(
            preprocessor,
            k_factors=1,
            factor_orders=2,
            compare_to=config['compare_to'],
            leading_prediction=config.get('leading_prediction', False),
            single_line=config.get('single_line', False),
            plot_y0=config.get('plot_y0', False),
        )
        modeler.run()

if SINGLE_BATCH_MODE == 'batch':
    industries = ['石油石化', '煤炭', '有色金属', '钢铁', '基础化工', '建材', '农林牧渔', '建筑',
                  '电子', '传媒', '计算机',
                  '机械', '电新', '电力', '国防军工',
                  '汽车', '家电', '医药', '食品饮料', '白酒', '食品', '饮料', '消费服务', '纺服', '商贸零售',
                  '交通运输',
                  '房地产', '非银金融']
    financial_list = ['净资产收益率ROE']  # , '归属母公司股东的净利润同比增长率', '营业收入同比增长率']

    def process_industry(industry, stationary, base_config, financial_list):
        preprocessor = DataPreprocessor(base_config, industry=industry, stationary=stationary)
        preprocessor.preprocess()

        results = []
        for financial in financial_list:
            modeler = DynamicFactorModeler(preprocessor, k_factors=1, factor_orders=2, compare_to=financial)
            modeler.apply_dynamic_factor_model()
            modeler.evaluate_model()

            # start_date = '2024-02-29'
            # end_date = '2024-04-30'
            start_date = None
            end_date = None

            contribution_text = f"{industry}行业 中观数据对综合景气指数的影响拆解({start_date} to {end_date}):\n"
            contribution_text += modeler.analyze_factor_contribution(start_date, end_date)

            results.append({
                'industry': industry,
                'stationary': stationary,
                'financial': financial,
                'figure': modeler.plot_factors(save_or_show='save'),
                'extracted_factor_filtered': modeler.results.factors.filtered['0'],
                'factor_filtered': modeler.series_compared_to.dropna().astype(float),
                'contribution_text': contribution_text
            })

        return results


    all_results = []
    for industry in industries:
        print(f'Processing industry {industry}')
        for stationary in [True, False]:
            results = process_industry(industry, stationary, base_config, financial_list)
            all_results.extend(results)

    document = Document()
    current_time = datetime.now().strftime("%Y%m%d_%H%M")
    file_path = rf'{base_config.excels_path}/景气/output_docs/各行业景气综合指数数据_{current_time}.xlsx'

    with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
        for result in all_results:
            stationary_str = '景气趋势慢线(中观数据平滑后)' if result['stationary'] else '景气变化快线(中观数据未平滑)'
            document.add_heading(f"{result['industry']} ({stationary_str}, financial={result['financial']})", level=1)
            document.add_picture(result['figure'], width=Inches(6))
            document.add_paragraph(result['contribution_text'])
            document.add_page_break()

            # 保存绘图数据到 Excel
            sheet_name = f"{result['industry']}"

            # 获取数据
            extracted_factor_filtered = result['extracted_factor_filtered']
            factor_filtered = result['factor_filtered']

            # 合并数据
            combined_df = pd.concat([extracted_factor_filtered, factor_filtered], axis=1)

            # 保存合并后的 DataFrame 到 Excel
            combined_df.to_excel(writer, sheet_name=sheet_name)

            # 获取工作簿和工作表
            workbook = writer.book
            worksheet = workbook.create_sheet(f"{sheet_name}_plot")

            # 将图片添加到工作表
            img = openpyxl.drawing.image.Image(result['figure'])
            worksheet.add_image(img, 'A1')

    document.save(rf'{base_config.excels_path}/景气/output_docs/各行业景气综合指数_{current_time}.docx')

if SINGLE_BATCH_MODE == 'batch_macro':
    macro_configs = [
        {'industry': '就业状况', 'compare_to': '中国:城镇调查失业率', 'stationary': False, 'date_start': '2018-01-01'},
        {'industry': '社零综指', 'compare_to': '中国:社会消费品零售总额:当月同比', 'stationary': False, 'date_start': '2020-01-01'},
        {'industry': '出口', 'compare_to': '中国:出口金额:当月同比', 'stationary': False, 'date_start': '2010-01-01'},
        {'industry': '出口', 'compare_to': '中国:出口金额:当月同比', 'stationary': False, 'date_start': '2010-01-01', 'single_line': True},
        {'industry': 'PPI', 'compare_to': '中国:PPI:全部工业品:当月同比', 'stationary': False, 'date_start': '2010-01-01', 'single_line': True},
        {'industry': '工业增加值', 'compare_to': '中国:规模以上工业增加值:当月同比', 'stationary': False, 'date_start': '2010-01-01'},
        {'industry': '制造业投资', 'compare_to': '(月度化)中国:固定资产投资完成额:制造业:累计同比', 'stationary': False, 'date_start': '2014-01-01', 'leading_prediction': False},
        {'industry': '房地产投资', 'compare_to': '(月度化)中国:房地产开发投资完成额:累计同比', 'stationary': False, 'date_start': '2012-01-01', 'leading_prediction': False, 'plot_y0': True},
        {'industry': '基建投资', 'compare_to': '(月度化)中国:固定资产投资完成额:基础设施建设投资:累计同比', 'stationary': False, 'date_start': '2012-01-01', 'leading_prediction': False, 'plot_y0': True},
    ]

    def process_macro_industry(config, base_config):
        preprocessor = DataPreprocessor(base_config, industry=config['industry'],
                                        stationary=config.get('stationary', False),
                                        date_start=config.get('date_start', '2013-01-01'),
                                        compare_to=config['compare_to']
                                        )
        preprocessor.preprocess()

        modeler = DynamicFactorModeler(
            preprocessor,
            k_factors=1,
            factor_orders=2,
            compare_to=config['compare_to'],
            leading_prediction=config.get('leading_prediction', False),
            single_line=config.get('single_line', False),
            plot_y0=config.get('plot_y0', False),
        )
        modeler.apply_dynamic_factor_model()
        modeler.evaluate_model()

        start_date = config.get('date_start', None)
        end_date = None

        contribution_text = f"{config['industry']}行业 中观数据对综合景气指数的影响拆解({start_date} to {end_date}):\n"
        contribution_text += modeler.analyze_factor_contribution(start_date, end_date)

        result = {
            'industry': config['industry'],
            'stationary': config['stationary'],
            'financial': config['compare_to'],
            'extracted_factor_filtered': modeler.results.factors.filtered['0'],
            'factor_filtered': modeler.series_compared_to.dropna().astype(float),
            'figure': modeler.plot_factors(save_or_show='save'),
            'contribution_text': contribution_text
        }

        return result


    all_results = []
    for config in macro_configs:
        print(f"Processing macro industry {config['industry']}")
        result = process_macro_industry(config, base_config)
        all_results.append(result)

    document = Document()
    current_time = datetime.now().strftime("%Y%m%d_%H%M")
    file_path = rf'{base_config.excels_path}/景气/output_docs/宏观经济景气综合指数数据_{current_time}.xlsx'

    with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
        for result in all_results:
            stationary_str = '景气趋势慢线(中观数据平滑后)' if result['stationary'] else '景气变化快线(中观数据未平滑)'
            document.add_heading(f"{result['industry']} ({stationary_str}, financial={result['financial']})", level=1)
            document.add_picture(result['figure'], width=Inches(6))
            document.add_paragraph(result['contribution_text'])
            document.add_page_break()

            # 保存绘图数据到 Excel
            sheet_name = f"{result['industry']}"

            # 获取数据
            extracted_factor_filtered = result['extracted_factor_filtered']
            factor_filtered = result['factor_filtered']

            # 合并数据
            combined_df = pd.concat([extracted_factor_filtered, factor_filtered], axis=1)

            # 保存合并后的 DataFrame 到 Excel
            combined_df.to_excel(writer, sheet_name=sheet_name)

            # 获取工作簿和工作表
            workbook = writer.book
            worksheet = workbook.create_sheet(f"{sheet_name}_plot")

            # 将图片添加到工作表
            img = openpyxl.drawing.image.Image(result['figure'])
            worksheet.add_image(img, 'A1')

    document.save(rf'{base_config.excels_path}/景气/output_docs/宏观经济景气综合指数_{current_time}.docx')